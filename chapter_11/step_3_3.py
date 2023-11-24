#######################################
# 1. 필요모듈
#######################################
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Mm, Pt

import step_0
import step_3_1
import step_3_2

#######################################
# 2. 환경설정
#######################################
STEP_3_3 = step_0.OUTPUT_FOLDER / "step_3_3.docx"


#######################################
# 3. 기본함수
#######################################
pass


#######################################
# 4. 메인함수
#######################################
def main():
    document = Document(step_3_2.STEP_3_2)

    p_style = document.styles["List Bullet"]
    p_format = p_style.paragraph_format
    p_format.space_before = 0
    p_format.space_after = 0
    p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    step_3_1.set_font_style(p_style, "나눔고딕", size=Pt(8))

    step_3_1.set_font_style(document.add_paragraph().add_run(" "), size=Pt(10))

    table = document.add_table(rows=1, cols=1)
    table.style = "Light Shading Accent 6"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    tr = table.rows[0]
    td = tr.cells[0]
    td.width = Mm(174)

    td.paragraphs[-1].add_run("주의사항")
    td.paragraphs[-1].paragraph_format.space_before = Mm(2)
    td.paragraphs[-1].paragraph_format.space_after = Mm(1)

    td.add_paragraph(style="List Bullet").text = "금융회사의 상품별 이자율 등 거래조건이 수시로 변경되어 지연공시될 수 있으므로 거래전 반드시 해당 금융회사에 문의하시기 바랍니다."
    td.paragraphs[-1].runs[-1].font.bold = False

    td.add_paragraph(style="List Bullet").text = "세전 이자율은 우대조건을 반영하지 않은 기본금리입니다. 상세정보의 우대조건에 해당시 보다 높은 이자율이 적용될 수 있습니다."
    td.paragraphs[-1].runs[-1].font.bold = False

    td.add_paragraph(style="List Bullet").text = "세후 이자율은 이자소득 원천징수세 15.4%(소득세 14%, 지방소득세 1.4%)를 차감한 금리입니다."
    td.paragraphs[-1].runs[-1].font.bold = False

    td.paragraphs[-1].paragraph_format.space_after = Mm(2)
    document.save(STEP_3_3)


#######################################
# 5. 실행
#######################################
if __name__ == "__main__":
    step_0.init_output_folder()
    main()

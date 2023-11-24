#######################################
# 1. 필요모듈
#######################################
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

import step_0
import step_1_2
import step_1_3

#######################################
# 2. 환경설정
#######################################
STEP_3_1 = step_0.OUTPUT_FOLDER / "step_3_1.docx"
PAGE_WIDTH, PAGE_HEIGHT, PAGE_MARGIN = Mm(210), Mm(297), Mm(12.7)
GRAPH_WIDTH, GRAPH_HEIGHT, CELL_WIDTH = Mm(30), Mm(8), Mm(35.5)


#######################################
# 3. 기본함수
#######################################
def set_page_margin(document):
    for section in document.sections:
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        section.top_margin = PAGE_MARGIN
        section.left_margin = PAGE_MARGIN
        section.right_margin = PAGE_MARGIN
        section.bottom_margin = PAGE_MARGIN


def set_font_style(obj, face=None, size=None, bold=None, color=None):
    if hasattr(obj, "font"):
        font = obj.font
        if face:
            font.name = face
            obj._element.rPr.rFonts.set(qn("w:eastAsia"), face)
        if size:
            font.size = size
        if bold:
            font.bold = bold
        if color:
            font.color.rgb = RGBColor.from_string(color)
        return obj


#######################################
# 4. 메인함수
#######################################
def main():
    document = Document()
    set_page_margin(document)

    p_title = document.add_paragraph("", style="Title")
    print(f"* add_run 실행 전 p_title.runs : {p_title.runs}")

    run_title = p_title.add_run("정기예금 금리 현황표")
    set_font_style(run_title, "나눔고딕", Pt(20), True, "0343DF")

    now_format = datetime.now().isoformat(sep=" ", timespec="minutes")
    run_datetime = p_title.add_run(f" ({now_format})")
    set_font_style(run_datetime, "나눔고딕", Pt(14), True, "0343df")

    style_normal = document.styles["Normal"]
    set_font_style(style_normal, face="나눔고딕", size=Pt(10))
    p_format = style_normal.paragraph_format
    p_format.space_before = 0
    p_format.space_after = 0
    p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p_blank = document.add_paragraph()
    set_font_style(p_blank.add_run(" "), size=Pt(6))

    p_1 = document.add_paragraph("")
    set_font_style(p_1.add_run("1. 주요 경제지표"), size=Pt(14), bold=True)

    p_blank = document.add_paragraph()
    set_font_style(p_blank.add_run(" "), size=Pt(10))

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    row_1 = table.rows[0]

    with pd.ExcelFile(step_1_2.STEP_1_2) as xlsx:
        sheet_names = [x for x in xlsx.sheet_names if x != "기준금리M"]
        for idx, name in enumerate(sheet_names):
            print(idx, name)
            df_raw = pd.read_excel(xlsx, sheet_name=name, index_col="TIME")
            s_value = df_raw["DATA_VALUE"]
            s_diff = s_value.diff()
            s_pct_change = s_value.pct_change()

            last_value = s_value.iloc[-1]
            last_diff = s_diff.iloc[-1]
            last_diff_rate = s_pct_change.iloc[-1]
            rgb = "FF0000" if last_diff > 0 else "0000FF" if last_diff < 0 else "000000"
            last_dt = s_value.index[-1].to_pydatetime()
            img_path = step_1_3.STEP_1_3.as_posix().format(
                f"{name}M" if name == "기준금리" else name
            )

            arrow = "▲" if last_diff > 0 else "▼" if last_diff < 0 else ""
            change = f"{arrow}{last_diff:,.2f}  {last_diff_rate:+,.2%}"

            cell = row_1.cells[idx]
            cell.width = CELL_WIDTH
            [cell.add_paragraph() for _ in range(4)]
            p1, p2, p3, p4, p5 = cell.paragraphs

            set_font_style(
                p1.add_run(name),
                size=Pt(12),
                bold=True,
                color="333333",
            )
            set_font_style(
                p2.add_run(f"{last_value:,.2f}"),
                size=Pt(14),
                bold=True,
                color="333333",
            )
            set_font_style(
                p3.add_run(change),
                size=Pt(10),
                bold=True,
                color=rgb,
            )

            p4.add_run().add_picture(img_path, width=GRAPH_WIDTH, height=GRAPH_HEIGHT)
            p4.paragraph_format.space_after = Mm(1)
            p4.paragraph_format.space_before = Mm(1)
            p4.paragraph_format.left_indent = Mm(-1)

            set_font_style(
                p5.add_run(last_dt.strftime("%Y-%m-%d")),
                size=Pt(8),
                bold=True,
                color="888888",
            )

    document.save(STEP_3_1)


#######################################
# 5. 실행
#######################################
if __name__ == "__main__":
    step_0.init_output_folder()
    main()

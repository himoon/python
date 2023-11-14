##############################################################################
# 1. 필요모듈
##############################################################################
import json
import pandas as pd

from docx import Document
from docx.shared import Inches, Cm, Pt, Mm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.enum.style import WD_STYLE


import matplotlib.pyplot as plt
import numpy as np
import time


##############################################################################
# 2. 환경설정
##############################################################################
STEP_1_2 = "output/step_1_2.xlsx"
STEP_1_3 = "output/{}"
STEP_3 = "output/step_3.docx"
TEMPLATE = "template.docx"


##############################################################################
# 3. 기본함수
##############################################################################
def set_margin(document, margin=12.7):
    for section in document.sections:
        section.top_margin = Mm(margin)
        section.bottom_margin = Mm(margin)
        section.left_margin = Mm(margin)
        section.right_margin = Mm(margin)
        section.page_width = Mm(210)
        section.page_height = Mm(297)


def init_style(document):
    style = document.styles["Normal"]
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # style = styles.add_style("normal_py", WD_STYLE_TYPE.PARAGRAPH)
    # style.base_style = styles["Normal"]


##############################################################################
# 4. 메인함수
##############################################################################
def main():
    document = Document()
    set_margin(document)
    init_style(document)

    # document.add_paragraph("하하하", style="normal_py")
    document.add_paragraph("하하하", style="No Spacing")
    document.add_paragraph("하하하", style="Normal")
    document.add_paragraph("", style="Normal")
    document.save(STEP_3)

    document.add_heading("정기예금 금리 현황표", 0)

    with pd.ExcelFile(STEP_1_2) as xlsx:
        xlsx.sheet_names
        df_bok = pd.read_excel(xlsx, sheet_name="bok")
        df_tb = pd.read_excel(xlsx, sheet_name="tb")
        df_cb = pd.read_excel(xlsx, sheet_name="cb")
        df_kospi = pd.read_excel(xlsx, sheet_name="kospi")
        df_ex = pd.read_excel(xlsx, sheet_name="ex")

    table = document.add_table(rows=1, cols=7)
    # table.add_row(style="No Spacing")
    # para0 = hdr_cells[0].add_paragraph(style="")
    # dir(para0.paragraph_format)
    # para0.paragraph_format
    # dir(para0)
    # dir(hdr_cells[0])
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "지표명"
    hdr_cells[1].text = "주기"
    hdr_cells[2].text = "값"
    hdr_cells[3].text = "단위"
    hdr_cells[4].text = "직전대비"
    hdr_cells[5].text = "등락률"
    hdr_cells[6].text = "그래프"
    document.save(STEP_3)

    s_last = df_kospi.iloc[-1]
    row_cells = table.add_row().cells
    row_cells[0].text = f'{s_last["ITEM_NAME1"]}'
    row_cells[1].text = "월"
    row_cells[2].text = f'{s_last["DATA_VALUE"]}'
    row_cells[3].text = f'{s_last["UNIT_NAME"]}'

    s_value = df_kospi["DATA_VALUE"]
    row_cells[4].text = f"{s_value.diff().iloc[-1]:+,.2f}"
    row_cells[5].text = f"{s_value.pct_change().iloc[-1]:+,.2%}"

    para = row_cells[6].paragraphs[0]
    run = para.add_run()
    pic = "output/grah_kospi.png"
    run.add_picture(pic, width=1400000, height=1400000)

    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.save(STEP_3)

    # p = document.add_paragraph("A plain paragraph having some ")
    # p.add_run("bold").bold = True
    # p.add_run(" and some ")
    # p.add_run("italic.").italic = True

    # document.add_heading("Heading, level 1", level=1)
    # document.add_paragraph("Intense quote", style="Intense Quote")

    # document.add_paragraph("first item in unordered list", style="List Bullet")
    # document.add_paragraph("first item in ordered list", style="List Number")

    # document.add_picture("test.png", width=Inches(1.25))

    records = (
        (3, "101", "Spam"),
        (7, "422", "Eggs"),
        (4, "631", "Spam, spam, eggs, and spam"),
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Qty"
    hdr_cells[1].text = "Id"
    hdr_cells[2].text = "Desc"
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    # document.add_page_break()
    # document.save("demo.docx")


##############################################################################
# 5. 실행
##############################################################################
if __name__ == "__main__":
    pass

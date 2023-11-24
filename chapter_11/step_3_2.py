#######################################
# 1. 필요모듈
#######################################
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

import step_0
import step_2_2
import step_3_1

#######################################
# 2. 환경설정
#######################################
STEP_3_2 = step_0.OUTPUT_FOLDER / "step_3_2.docx"
NUM_OF_ROWS = 13


#######################################
# 3. 기본함수
#######################################
pass


#######################################
# 4. 메인함수
#######################################
def main():
    with pd.ExcelFile(step_2_2.STEP_2_2) as xlsx:
        print(f"sheet_names: {xlsx.sheet_names!r}")
        df_raw: pd.DataFrame = pd.read_excel(xlsx, sheet_name="deposit")
        df_filtered = df_raw.filter(
            ["금융회사", "상품명", "가입제한여부", "세전이자율", "세후이자율", "최고우대금리"]
        )

    document = Document(step_3_1.STEP_3_1)
    p_head = document.add_paragraph()
    p_head.paragraph_format.space_before = Mm(10)
    p_head.paragraph_format.space_after = Mm(2)
    step_3_1.set_font_style(p_head.add_run("2. 주요 은행 정기예금 금리"), size=Pt(14), bold=True)

    table = document.add_table(rows=1, cols=6)
    table.style = "Light Shading Accent 4"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    th = table.rows[0]
    th_text = ["금융회사", "상품명", "가입제한", "세전", "세후", "최고우대"]
    th_width = [Mm(40), Mm(53), Mm(20), Mm(20), Mm(20), Mm(20)]
    for idx, td in enumerate(th.cells):
        td.text = f"{th_text[idx]}"
        td.width = th_width[idx]
        td.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        td.paragraphs[-1].runs[-1].font.size = Pt(12)
        td.paragraphs[-1].runs[-1].font.bold = True
        td.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for _, s_row in df_filtered.head(NUM_OF_ROWS).iterrows():
        tr = table.add_row()
        for idx, td in enumerate(tr.cells):
            td = tr.cells[idx]
            td.text = f"{s_row.iloc[idx]}"
            td.width = th_width[idx]
            if idx < 2:
                td.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
            td.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            td.paragraphs[-1].paragraph_format.space_after = Mm(2)
            td.paragraphs[-1].paragraph_format.space_before = Mm(2)

    document.save(STEP_3_2)


#######################################
# 5. 실행
#######################################
if __name__ == "__main__":
    step_0.init_output_folder()
    main()

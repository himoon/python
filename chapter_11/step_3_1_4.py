#######################################
# 1. 필요모듈
#######################################
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

import step_0
import step_1_2
import step_1_3

#######################################
# 2. 환경설정
#######################################
STEP_3_1 = step_0.OUTPUT_FOLDER / "step_3_1.docx"
PAGE_WIDTH, PAGE_HEIGHT, PAGE_MARGIN = Mm(210), Mm(297), Mm(12.7)
GRAPH_WIDTH, GRAPH_HEIGHT, CELL_WIDTH = Mm(30), Mm(8), Mm(35.5)


#######################################
# 3. 기본함수
#######################################
pass


#######################################
# 4. 메인함수
#######################################
def main():
    document = Document()

    for section in document.sections:
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        section.top_margin = PAGE_MARGIN
        section.left_margin = PAGE_MARGIN
        section.right_margin = PAGE_MARGIN
        section.bottom_margin = PAGE_MARGIN

    p_title = document.add_paragraph("", style="Title")
    print(f"* add_run 실행 전 p_title.runs : {p_title.runs}")

    run_title = p_title.add_run("정기예금 금리 현황표")
    run_title.font.bold = True
    run_title.font.size = Pt(20)
    run_title.font.color.rgb = RGBColor(0x03, 0x43, 0xDF)
    run_title.font.name = "나눔고딕"
    run_title._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")

    now_format = datetime.now().isoformat(sep=" ", timespec="minutes")
    run_datetime = p_title.add_run(f" ({now_format})")
    run_datetime.font.bold = True
    run_datetime.font.size = Pt(14)
    run_datetime.font.color.rgb = RGBColor.from_string("0343df")
    run_datetime.font.name = "나눔고딕"
    run_datetime._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")

    style_normal = document.styles["Normal"]
    style_normal.font.size = Pt(10)
    style_normal.font.name = "나눔고딕"
    style_normal._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
    p_format = style_normal.paragraph_format
    p_format.space_before = 0
    p_format.space_after = 0
    p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p_blank = document.add_paragraph()
    p_blank.add_run(" ").font.size = Pt(6)

    p_1 = document.add_paragraph("")
    run_1 = p_1.add_run("1. 주요 경제지표")
    run_1.font.bold = True
    run_1.font.size = Pt(14)

    p_blank = document.add_paragraph()
    p_blank.add_run(" ").font.size = Pt(10)

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    row_1 = table.rows[0]

    with pd.ExcelFile(step_1_2.STEP_1_2) as xlsx:
        sheet_names = [x for x in xlsx.sheet_names if x != "기준금리M"]
        for idx, name in enumerate(sheet_names):
            print(idx, name)
            df_raw = pd.read_excel(xlsx, sheet_name=name, index_col="TIME")
            s_value = df_raw["DATA_VALUE"]
            s_diff = s_value.diff()
            s_pct_change = s_value.pct_change()

            last_value = s_value.iloc[-1]
            last_diff = s_diff.iloc[-1]
            last_diff_rate = s_pct_change.iloc[-1]
            rgb = "FF0000" if last_diff > 0 else "0000FF" if last_diff < 0 else "000000"
            last_dt = s_value.index[-1].to_pydatetime()
            img_path = step_1_3.STEP_1_3.as_posix().format(name)

            arrow = "▲" if last_diff > 0 else "▼" if last_diff < 0 else ""
            change = f"{arrow}{last_diff:,.2f}  {last_diff_rate:+,.2%}"

            cell = row_1.cells[idx]
            cell.width = CELL_WIDTH
            [cell.add_paragraph() for _ in range(4)]
            p1, p2, p3, p4, p5 = cell.paragraphs

            p1_run = p1.add_run(name)
            p1_run.font.size = Pt(12)
            p1_run.font.bold = True
            p1_run.font.color.rgb = RGBColor.from_string("333333")

            p2_run = p2.add_run(f"{last_value:,.2f}")
            p2_run.font.size = Pt(14)
            p2_run.font.bold = True
            p2_run.font.color.rgb = RGBColor.from_string("333333")

            p3_run = p3.add_run(change)
            p3_run.font.size = Pt(10)
            p3_run.font.bold = True
            p3_run.font.color.rgb = RGBColor.from_string(rgb)

            p4.add_run().add_picture(img_path, width=GRAPH_WIDTH, height=GRAPH_HEIGHT)
            p4.paragraph_format.space_after = Mm(1)
            p4.paragraph_format.space_before = Mm(1)
            p4.paragraph_format.left_indent = Mm(-1)

            p5_run = p5.add_run(last_dt.strftime("%Y-%m-%d"))
            p5_run.font.size = Pt(8)
            p5_run.font.bold = True
            p5_run.font.color.rgb = RGBColor.from_string("888888")

    document.save(STEP_3_1)


#######################################
# 5. 실행
#######################################
if __name__ == "__main__":
    step_0.init_output_folder()
    main()

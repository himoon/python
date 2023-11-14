##############################################################################
# 1. 필요모듈
##############################################################################
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

##############################################################################
# 2. 환경설정
##############################################################################
STEP_1_2 = "output/step_1_2.xlsx"
STEP_1_3 = "output/{}"
STEP_3_1 = "output/step_3_1.docx"
GRAPH_WIDTH, GRAPH_HEIGHT = Mm(30), Mm(8)


##############################################################################
# 3. 기본함수
##############################################################################
def set_margin(document, margin=12.7):
    for section in document.sections:
        section.top_margin = Mm(margin)
        section.bottom_margin = Mm(margin)
        section.left_margin = Mm(margin)
        section.right_margin = Mm(margin)
        section.page_width = Mm(210)
        section.page_height = Mm(297)


def apply_font_style(style, size=None, bold=False, font_name=None):
    this_font = style.font
    if size:
        this_font.size = size
    if bold:
        this_font.bold = True
    if font_name:
        this_font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    return style


def init_style(document):
    style_normal = document.styles["Normal"]
    apply_font_style(style_normal, size=Pt(10), font_name="나눔고딕")

    p_format = style_normal.paragraph_format
    p_format.space_before = 0
    p_format.space_after = 0
    p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def get_index_values(df_raw):
    s_value = df_raw["DATA_VALUE"]
    s_diff = s_value.diff()
    s_pct_change = s_value.pct_change()

    value = s_value.iloc[-1]
    diff = s_diff.iloc[-1]
    diff_rate = s_pct_change.iloc[-1]
    dt = s_value.index[-1].to_pydatetime()
    return value, diff, diff_rate, dt


def get_arrow(diff):
    return "▲" if diff > 0 else "▼" if diff < 0 else ""


def add_break_line(document, pt):
    document.add_paragraph().style.font.size = Pt(pt)


##############################################################################
# 4. 메인함수
##############################################################################
def main():
    document = Document()
    set_margin(document)
    init_style(document)

    ##########################################################################
    ##########################################################################
    p_title = document.add_heading("정기예금 금리 현황표", level=0)
    apply_font_style(p_title.runs[-1], Pt(20), True, "나눔고딕")

    p_title.add_run(datetime.now().strftime(f"(%Y.%m.%d.)"))
    apply_font_style(p_title.runs[-1], Pt(14), True, "나눔고딕")

    add_break_line(document, 6)

    p_head_1 = document.add_paragraph()
    apply_font_style(p_head_1.add_run("1. 주요 경제지표"), Pt(14), True)

    add_break_line(document, 6)
    document.save(STEP_3_1)

    ##########################################################################
    ##########################################################################
    with pd.ExcelFile(STEP_1_2) as xlsx:
        xlsx.sheet_names
        df_base: pd.DataFrame = pd.read_excel(xlsx, sheet_name="base", index_col="TIME")
        df_tb: pd.DataFrame = pd.read_excel(xlsx, sheet_name="tb", index_col="TIME")
        df_cb: pd.DataFrame = pd.read_excel(xlsx, sheet_name="cb", index_col="TIME")
        df_kospi: pd.DataFrame = pd.read_excel(
            xlsx, sheet_name="kospi", index_col="TIME"
        )
        df_ex: pd.DataFrame = pd.read_excel(xlsx, sheet_name="ex", index_col="TIME")

    ##########################################################################
    ##########################################################################
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row_1_1 = table.rows[0]
    cell_a, cell_b, cell_c, cell_d, cell_e = row_1_1.cells

    ##########################################################################
    # 기준금리
    ##########################################################################
    value, diff, diff_rate, dt = get_index_values(df_base)
    change = f"{get_arrow(diff)}{diff:,.2f}  {diff_rate:+,.2%}"

    p1 = cell_a.paragraphs[0]
    p2, p3, p4, p5 = [cell_a.add_paragraph() for _ in range(4)]
    apply_font_style(p1.add_run("기준금리"), Pt(14), True)
    apply_font_style(p2.add_run(f"{value:,.2f}"), Pt(20), True)
    apply_font_style(p3.add_run(change), Pt(10), True)  # diff, diff_rate
    apply_font_style(p5.add_run(dt.strftime("%Y.%m.")), Pt(8), True)
    p4.add_run().add_picture(
        f"output/grah_base_mo.png", width=GRAPH_WIDTH, height=GRAPH_HEIGHT
    )
    document.save(STEP_3_1)

    ##########################################################################
    # 국고채
    ##########################################################################
    value, diff, diff_rate, dt = get_index_values(df_tb)
    change = f"{get_arrow(diff)}{diff:,.3f}  {diff_rate:+,.2%}"

    p1 = cell_b.paragraphs[0]
    p2, p3, p4, p5 = [cell_b.add_paragraph() for _ in range(4)]

    apply_font_style(p1.add_run("국고채"), Pt(14), True)
    apply_font_style(p1.add_run("(3Y)"), Pt(8), True)
    apply_font_style(p2.add_run(f"{value:,.3f}"), Pt(20), True)
    apply_font_style(p3.add_run(change), Pt(10), True)
    apply_font_style(p5.add_run(dt.strftime("%Y.%m.%d.")), Pt(8), True)
    p4.add_run().add_picture(
        f"output/grah_tb.png", width=GRAPH_WIDTH, height=GRAPH_HEIGHT
    )
    document.save(STEP_3_1)

    ##########################################################################
    # 회사채
    ##########################################################################
    value, diff, diff_rate, dt = get_index_values(df_cb)
    change = f"{get_arrow(diff)}{diff:,.3f}  {diff_rate:+,.2%}"

    p1 = cell_c.paragraphs[0]
    p2, p3, p4, p5 = [cell_c.add_paragraph() for _ in range(4)]

    apply_font_style(p1.add_run("회사채"), Pt(14), True)
    apply_font_style(p1.add_run("(3Y,AA-)"), Pt(8), True)
    apply_font_style(p2.add_run(f"{value:,.3f}"), Pt(20), True)
    apply_font_style(p3.add_run(change), Pt(10), True)
    apply_font_style(p5.add_run(dt.strftime("%Y.%m.%d.")), Pt(8), True)
    p4.add_run().add_picture(
        f"output/grah_cb.png", width=GRAPH_WIDTH, height=GRAPH_HEIGHT
    )
    document.save(STEP_3_1)

    ##########################################################################
    # KOSPI
    ##########################################################################
    value, diff, diff_rate, dt = get_index_values(df_kospi)
    change = f"{get_arrow(diff)}{diff:,.2f}  {diff_rate:+,.2%}"

    p1 = cell_d.paragraphs[0]
    p2, p3, p4, p5 = [cell_d.add_paragraph() for _ in range(4)]

    apply_font_style(p1.add_run("KOSPI"), Pt(14), True)
    apply_font_style(p2.add_run(f"{value:,.2f}"), Pt(20), True)
    apply_font_style(p3.add_run(change), Pt(10), True)
    apply_font_style(p5.add_run(dt.strftime("%Y.%m.%d.")), Pt(8), True)
    p4.add_run().add_picture(
        f"output/grah_kospi.png", width=GRAPH_WIDTH, height=GRAPH_HEIGHT
    )
    document.save(STEP_3_1)

    ##########################################################################
    # 원달러
    ##########################################################################
    value, diff, diff_rate, dt = get_index_values(df_ex)
    change = f"{get_arrow(diff)}{diff:,.2f}  {diff_rate:+,.2%}"

    p1 = cell_e.paragraphs[0]
    p2, p3, p4, p5 = [cell_e.add_paragraph() for _ in range(4)]

    apply_font_style(p1.add_run("원/달러환율"), Pt(14), True)
    apply_font_style(p2.add_run(f"{value:,.2f}"), Pt(20), True)
    apply_font_style(p3.add_run(change), Pt(10), True)
    apply_font_style(p5.add_run(dt.strftime("%Y.%m.%d.")), Pt(8), True)
    p4.add_run().add_picture(
        f"output/grah_ex.png", width=GRAPH_WIDTH, height=GRAPH_HEIGHT
    )
    document.save(STEP_3_1)


##############################################################################
# 5. 실행
##############################################################################
if __name__ == "__main__":
    main()
